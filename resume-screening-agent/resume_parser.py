"""Extract normalized text from .docx resumes (paragraphs + tables)."""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

_NAME_BLOCKLIST_PATTERNS = [
    re.compile(r"@"),
    re.compile(r"\d{3,}"),
    re.compile(r"^(resume|curriculum vitae|cv|profile|summary)\b", re.IGNORECASE),
    re.compile(r"^(email|phone|mobile|address|location|linkedin|github|www\.)", re.IGNORECASE),
]

_LOCATION_LABEL_PATTERN = re.compile(
    r"^(location|address|city|based in)\s*[:\-]\s*(.+)$", re.IGNORECASE
)

# Skills are often listed per-project as "Environment: X, Y, Z" rather than
# under one dedicated "Skills:" section, so match both.
_SKILLS_LABEL_PATTERN = re.compile(
    r"^(skills|technical skills|core competencies|key skills|technology|technologies|environment)"
    r"\s*[:\-|]\s*(.+)$",
    re.IGNORECASE,
)
_SKILLS_MAX_LENGTH = 400


@dataclass
class ParsedResume:
    filename: str
    filepath: str
    paragraphs: list[str] = field(default_factory=list)
    text: str = ""
    candidate_name: str = ""
    location: str = ""
    skills: str = ""


def _extract_table_text(document: Document) -> list[str]:
    lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cell_texts = []
            prev = None
            for cell in row.cells:
                cell_text = cell.text.strip()
                # merged cells repeat the same text in consecutive columns; skip repeats
                if cell_text and cell_text != prev:
                    cell_texts.append(cell_text)
                prev = cell_text

            if not cell_texts:
                continue

            if len(cell_texts) > 1 and all("\n" not in c for c in cell_texts):
                # simple "label | value" row (each cell is a single line)
                lines.append(" | ".join(cell_texts))
            else:
                # a cell holds an entire multi-paragraph section as one blob
                # (cell.text joins internal paragraphs with "\n") — split it
                # back into individual lines so line-anchored label matching
                # (e.g. "Environment: ...") still works
                for cell_text in cell_texts:
                    lines.extend(line for line in cell_text.splitlines() if line.strip())
    return lines


def _guess_candidate_name(all_lines: list[str], filename: str) -> str:
    for line in all_lines[:10]:
        candidate = line.strip()
        if not candidate:
            continue
        if len(candidate) > 60:
            continue
        word_count = len(candidate.split())
        if not (1 <= word_count <= 5):
            continue
        if any(pattern.search(candidate) for pattern in _NAME_BLOCKLIST_PATTERNS):
            continue
        return candidate
    return Path(filename).stem.replace("_", " ").replace("-", " ").strip().title()


def _guess_location(all_lines: list[str]) -> str:
    for line in all_lines[:60]:
        match = _LOCATION_LABEL_PATTERN.match(line.strip())
        if match:
            return match.group(2).strip()
    return ""


def _guess_skills(all_lines: list[str]) -> str:
    seen: set[str] = set()
    matches: list[str] = []
    for line in all_lines:
        match = _SKILLS_LABEL_PATTERN.match(line.strip())
        if not match:
            continue
        value = match.group(2).strip()
        if value and value not in seen:
            seen.add(value)
            matches.append(value)

    skills = "; ".join(matches)
    if len(skills) > _SKILLS_MAX_LENGTH:
        skills = skills[:_SKILLS_MAX_LENGTH].rsplit(" ", 1)[0] + "..."
    return skills


def parse_docx(filepath: Path) -> ParsedResume | None:
    try:
        document = Document(str(filepath))
    except Exception as exc:  # corrupt/unreadable file
        logger.warning("Skipping unreadable resume %s: %s", filepath.name, exc)
        return None

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    table_lines = _extract_table_text(document)
    all_lines = paragraphs + table_lines

    full_text = "\n".join(all_lines)
    if not full_text.strip():
        logger.warning("Skipping empty resume %s", filepath.name)
        return None

    return ParsedResume(
        filename=filepath.name,
        filepath=str(filepath.resolve()),
        paragraphs=paragraphs,
        text=full_text,
        candidate_name=_guess_candidate_name(all_lines, filepath.name),
        location=_guess_location(all_lines),
        skills=_guess_skills(all_lines),
    )


def parse_resume_folder(folder_path: str | Path) -> list[ParsedResume]:
    folder = Path(folder_path)
    if not folder.is_dir():
        raise FileNotFoundError(f"Resume folder not found: {folder}")

    docx_files = sorted(folder.glob("*.docx"))
    logger.info("Found %d .docx file(s) in %s", len(docx_files), folder)

    parsed: list[ParsedResume] = []
    skipped = 0
    for filepath in docx_files:
        if filepath.name.startswith("~$"):
            continue  # Word lock/temp file
        resume = parse_docx(filepath)
        if resume is None:
            skipped += 1
            continue
        parsed.append(resume)

    logger.info("Parsed %d resume(s), skipped %d", len(parsed), skipped)
    return parsed
